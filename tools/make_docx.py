# -*- coding: utf-8 -*-
"""
Builds ICNGCI-2027-Conference-Information.docx from the live site content.

Content is extracted from the HTML pages rather than retyped, so the document
and the website cannot drift apart. Re-run after editing the site.
"""
import os
import re
import html as H
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "assets", "downloads", "ICNGCI-2027-Conference-Information.docx")

INK = RGBColor(0x0A, 0x0A, 0x0A)
BLUE = RGBColor(0x00, 0x3B, 0xCE)
RED = RGBColor(0xFF, 0x2B, 0x37)
GREY = RGBColor(0x61, 0x61, 0x61)

# --------------------------------------------------------------------------
# HTML helpers
# --------------------------------------------------------------------------
def read(page):
    return open(os.path.join(ROOT, page), encoding="utf-8").read()


def text_of(fragment):
    """Strip tags, unescape entities, collapse whitespace."""
    t = re.sub(r"(?s)<(script|style)\b.*?</\1>", " ", fragment)
    t = re.sub(r"<br\s*/?>", " ", t)
    t = re.sub(r"(?s)<[^>]+>", " ", t)
    t = H.unescape(t)
    t = t.replace(" ", " ").replace("­", "")
    return re.sub(r"\s+", " ", t).strip()


def main_of(page):
    s = read(page)
    m = re.search(r'(?s)<main id="main">(.*?)</main>', s)
    return m.group(1) if m else s


def blocks(fragment, tags=("h2", "h3", "h4", "p", "li", "caption")):
    """Yield (tag, text) in document order."""
    pat = re.compile(r"<(%s)\b[^>]*>(.*?)</\1>" % "|".join(tags), re.S | re.I)
    for m in pat.finditer(fragment):
        t = text_of(m.group(2))
        if t:
            yield m.group(1).lower(), t


def tables_of(fragment):
    """Extract <table class="data"> as (caption, headers, rows)."""
    out = []
    for tm in re.finditer(r'(?s)<table class="data">(.*?)</table>', fragment):
        body = tm.group(1)
        cap = re.search(r"(?s)<caption[^>]*>(.*?)</caption>", body)
        caption = text_of(cap.group(1)) if cap else ""
        headers = [text_of(c) for c in re.findall(r"(?s)<th[^>]*>(.*?)</th>", body)]
        rows = []
        for rm in re.finditer(r"(?s)<tr>(.*?)</tr>", body):
            cells = re.findall(r"(?s)<t[dh][^>]*>(.*?)</t[dh]>", rm.group(1))
            vals = [text_of(c) for c in cells]
            if vals and not all(v in headers for v in vals):
                rows.append(vals)
        out.append((caption, headers, rows))
    return out


# --------------------------------------------------------------------------
# Document helpers
# --------------------------------------------------------------------------
doc = Document()

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.15

for sec in doc.sections:
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)


def para(text, size=10.5, bold=False, color=None, space_before=0, space_after=6,
         align=None, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    return p


def h1(text):
    doc.add_page_break()
    para(text.upper(), size=8, bold=True, color=BLUE, space_after=2)
    p = para(text, size=20, bold=True, color=INK, space_after=10)
    return p


def h2(text):
    para(text, size=13.5, bold=True, color=INK, space_before=14, space_after=5)


def h3(text):
    para(text, size=11, bold=True, color=INK, space_before=10, space_after=3)


def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.size = Pt(10.5)
    return p


def numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.size = Pt(10.5)
    return p


def table(headers, rows, widths=None):
    if not rows:
        return
    ncol = max(len(headers) if headers else 0, max(len(r) for r in rows))
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    if headers:
        cells = t.add_row().cells
        for i, htxt in enumerate(headers[:ncol]):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(htxt)
            r.bold = True
            r.font.size = Pt(9)
            shade = OxmlElement("w:shd")
            shade.set(qn("w:fill"), "0A0A0A")
            cells[i]._tc.get_or_add_tcPr().append(shade)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row[:ncol]):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(v)
            r.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths[:ncol]):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ==========================================================================
# TITLE PAGE
# ==========================================================================
para("Sharda School of Computing Science & Engineering", size=9, bold=True,
     color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("SHARDA UNIVERSITY, GREATER NOIDA", size=9, bold=True, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

para("ICNGCI 2027", size=40, bold=True, color=INK,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
para("International Conference on\nNext-Generation Computing and Innovations",
     size=16, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
para("17–19 June 2027", size=14, bold=True, color=RED,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Sharda University, Knowledge Park III, Greater Noida, Uttar Pradesh 201310, India",
     size=10.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

para("Publication partner: Springer", size=11, bold=True, color=INK,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Submission portal: Microsoft CMT  ·  Paper deadline: 15 April 2027",
     size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

para("CONFERENCE INFORMATION DOCUMENT", size=8, bold=True, color=BLUE,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Complete content of the ICNGCI 2027 conference website.",
     size=9.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Generated from the site — re-run tools/make-docx.sh after editing pages.",
     size=8, color=GREY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ==========================================================================
# CONTENTS
# ==========================================================================
h1("Contents")
for i, name in enumerate([
        "About the Conference", "Technical Tracks", "Call for Papers",
        "Important Dates", "Paper Submission", "Registration and Fees",
        "Conference Committee", "Speakers", "Programme",
        "Venue, Travel and Accommodation", "Contact"], 1):
    para("%d.  %s" % (i, name), size=11, space_after=4)

# ==========================================================================
# 1. ABOUT
# ==========================================================================
h1("1. About the Conference")
frag = main_of("about.html")
skip_heads = {"On this page"}
current = None
for tag, txt in blocks(frag):
    if txt in skip_heads or txt in ("Contents",):
        current = "skip"
        continue
    if tag == "h2":
        current = txt
        h2(txt)
    elif tag == "h3":
        if current == "skip":
            continue
        h3(txt)
    elif tag == "li":
        if current == "skip":
            continue
        bullet(txt)
    elif tag == "p":
        if current == "skip":
            continue
        if len(txt) < 3:
            continue
        para(txt)

# ==========================================================================
# 2. TRACKS
# ==========================================================================
h1("2. Technical Tracks")
para("Six tracks, ninety topic areas. Submit to the track that best matches the "
     "primary contribution of the paper; the Technical Program Chairs reassign "
     "where a better fit exists.", space_after=10)

tsrc = read("tracks.html")
for m in re.finditer(r'(?s)<article class="track"[^>]*>(.*?)</article>', tsrc):
    body = m.group(1)
    num = text_of(re.search(r'(?s)<p class="track__num">(.*?)</p>', body).group(1))
    title = text_of(re.search(r"(?s)<h3>(.*?)</h3>", body).group(1))
    desc_m = re.search(r'(?s)<div class="track__head">.*?<p>(.*?)</p>\s*</div>', body)
    h2("%s — %s" % (num, title))
    if desc_m:
        para(text_of(desc_m.group(1)), color=GREY, space_after=5)
    topics = [text_of(t) for t in re.findall(r"(?s)<li[^>]*>(.*?)</li>", body)]
    for i, t in enumerate(topics, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("%2d.  " % i)
        r.font.size = Pt(10)
        r.font.color.rgb = BLUE
        r.bold = True
        r2 = p.add_run(t)
        r2.font.size = Pt(10)

# ==========================================================================
# 3. CALL FOR PAPERS
# ==========================================================================
h1("3. Call for Papers")
para("FULL PAPER SUBMISSION DEADLINE", size=8, bold=True, color=BLUE, space_after=2)
para("15 April 2027", size=22, bold=True, color=INK, space_after=2)
para("23:59 Anywhere on Earth · 8–10 pages, A4, Springer template · Microsoft CMT",
     size=9.5, color=GREY, space_after=12)

frag = main_of("call-for-papers.html")
frag = re.sub(r'(?s)<div class="deadline">.*?</div>\s*</div>\s*</section>', "", frag, count=1)
seen_tracks_glance = False
for tag, txt in blocks(frag):
    if txt in ("Key facts", "Downloads", "Tracks at a glance", "Scope"):
        seen_tracks_glance = True
        continue
    if seen_tracks_glance:
        continue
    if tag == "h2":
        h2(txt)
    elif tag == "h3":
        h3(txt)
    elif tag == "li":
        bullet(txt)
    elif tag == "p" and len(txt) > 2:
        para(txt)
for cap, hd, rows in tables_of(main_of("call-for-papers.html")):
    if cap:
        h3(cap)
    table(hd, rows)

# ==========================================================================
# 4. IMPORTANT DATES
# ==========================================================================
h1("4. Important Dates")
rows = []
for m in re.finditer(r'(?s)<li data-date="([^"]+)"[^>]*>(.*?)</li>', read("dates.html")):
    label = text_of(re.search(r'(?s)dates__label">(.*?)</span>', m.group(2)).group(1))
    when = text_of(re.search(r'(?s)dates__when">(.*?)</span>', m.group(2)).group(1))
    rows.append([label, when])
table(["Milestone", "Date"], rows, widths=[4.3, 2.3])
para("All deadlines close at 23:59 Anywhere on Earth (AoE). Any extension is "
     "announced on the conference website and by email to registered authors.",
     size=9.5, color=GREY, space_before=4)

for cap, hd, tbl in tables_of(main_of("dates.html")):
    h3("What each deadline requires")
    table(hd, tbl, widths=[1.5, 3.4, 1.7])
    break

# ==========================================================================
# 5. PAPER SUBMISSION
# ==========================================================================
h1("5. Paper Submission")
frag = main_of("submission.html")
frag = re.sub(r"(?s)<aside>.*?</aside>", "", frag)
for tag, txt in blocks(frag):
    if txt in ("Downloads", "Key facts"):
        continue
    if tag == "h2":
        h2(txt)
    elif tag == "h3":
        h3(txt)
    elif tag == "li":
        bullet(txt)
    elif tag == "p" and len(txt) > 2:
        para(txt)

# ==========================================================================
# 6. REGISTRATION
# ==========================================================================
h1("6. Registration and Fees")
para("One registration covers one accepted paper. At least one author of every "
     "accepted paper must register and present the work — on-site or online — "
     "for the paper to appear in the proceedings.", space_after=10)
for cap, hd, rows in tables_of(main_of("registration.html")):
    if cap:
        h3(cap)
    table(hd, rows, widths=[2.2, 1.5, 1.5, 1.4])

h2("What registration includes")
inc = re.search(r"(?s)<h3>What registration includes</h3>(.*?)</ul>", read("registration.html"))
if inc:
    for t in re.findall(r"(?s)<li>(.*?)</li>", inc.group(1)):
        bullet(text_of(t))

h2("Payment")
kv = re.search(r"(?s)<h3>Bank account details</h3>(.*?)</ul>", read("registration.html"))
if kv:
    rows = []
    for li in re.findall(r"(?s)<li>(.*?)</li>", kv.group(1)):
        k = text_of(re.search(r'(?s)kv__k">(.*?)</span>', li).group(1))
        v = text_of(re.sub(r"(?s)<button.*?</button>", "",
                           re.search(r'(?s)kv__v[^"]*">(.*?)</span>', li).group(1)))
        rows.append([k, v])
    table(["Field", "Value"], rows, widths=[1.8, 4.8])
para("Verify every field against a current document from the university finance "
     "office before circulating this document.", size=9.5, color=RED, italic=True)

h2("Policies")
pol = re.search(r"(?s)<h3>Policies</h3>(.*?)</div>", read("registration.html"))
if pol:
    for tag, txt in blocks(pol.group(1), tags=("h4", "p")):
        if tag == "h4":
            h3(txt)
        else:
            para(txt)

# ==========================================================================
# 7. COMMITTEE
# ==========================================================================
h1("7. Conference Committee")
csrc = read("committee.html")
panels = {
    "patrons": "Patrons and Chairs",
    "advisory": "Advisory Committee",
    "tpc": "Technical Program Committee",
    "organizing": "Organizing Committee",
}
for pid, ptitle in panels.items():
    start = csrc.find('<div id="%s" role="tabpanel"' % pid)
    if start == -1:
        continue
    nxt = min([i for i in
               (csrc.find('<div id="%s" role="tabpanel"' % o, start + 1) for o in panels)
               if i > start] or [len(csrc)])
    panel = csrc[start:nxt]
    if pid == "organizing":
        panel = panel.split('<div class="note"')[0]
    h2(ptitle)
    for gm in re.finditer(r"(?s)<div data-people-group>(.*?)</div>\n      </div>", panel):
        g = gm.group(1)
        gh = re.search(r"(?s)<h3[^>]*>(.*?)</h3>", g)
        if gh:
            h3(text_of(gh.group(1)))
        rows = []
        for pm in re.finditer(r'(?s)<div class="person">(.*?)</div>\s*</div>', g):
            b = pm.group(1)
            role = re.search(r'(?s)person__role">(.*?)</p>', b)
            name = re.search(r'(?s)person__name">(.*?)</p>', b)
            aff = re.search(r'(?s)person__affil">(.*?)</p>', b)
            if name:
                rows.append([text_of(name.group(1)),
                             text_of(role.group(1)) if role else "",
                             text_of(aff.group(1)) if aff else ""])
        table(["Name", "Role", "Affiliation"], rows, widths=[1.9, 1.4, 3.3])
    # track chairs / free-standing person cards outside groups
    if pid == "tpc":
        h3("Track chairs")
        rows = []
        for pm in re.finditer(r'(?s)<div class="person">(.*?)</div>\s*</div>', panel):
            b = pm.group(1)
            role = re.search(r'(?s)person__role">(.*?)</p>', b)
            name = re.search(r'(?s)person__name">(.*?)</p>', b)
            aff = re.search(r'(?s)person__affil">(.*?)</p>', b)
            if name:
                rows.append([text_of(name.group(1)),
                             text_of(role.group(1)) if role else "",
                             text_of(aff.group(1)) if aff else ""])
        table(["Name", "Role", "Track / department"], rows, widths=[1.9, 1.3, 3.4])
        h3("Technical Program Committee members")
        tpc_rows = []
        tm = re.search(r'(?s)<h3[^>]*>Technical Program Committee members</h3>.*?<tbody>(.*?)</tbody>', panel)
        if tm:
            for rm in re.finditer(r"(?s)<tr>(.*?)</tr>", tm.group(1)):
                cells = [text_of(c) for c in re.findall(r"(?s)<td[^>]*>(.*?)</td>", rm.group(1))]
                if len(cells) == 2:
                    tpc_rows.append(cells)
        para("%d reviewers." % len(tpc_rows), size=9.5, color=GREY, space_after=4)
        table(["Name", "Affiliation"], tpc_rows, widths=[2.4, 4.2])

# ==========================================================================
# 8. SPEAKERS
# ==========================================================================
h1("8. Speakers")
para("The line-up carries forward from the previous conference hosted at Sharda "
     "University and is being reconfirmed for ICNGCI 2027. Talk titles are indicative.",
     color=GREY, space_after=10)
ssrc = read("speakers.html")
for sec in re.finditer(r'(?s)<div><span class="eyebrow">([^<]*)</span><h2>([^<]*)</h2></div>(.*?)(?=<div class="section-head"|</section>)', ssrc):
    h2(sec.group(2))
    rows = []
    for pm in re.finditer(r'(?s)<article class="person person--feature"[^>]*>(.*?)</article>', sec.group(3)):
        b = pm.group(1)
        role = re.search(r'(?s)person__role">(.*?)</p>', b)
        name = re.search(r'(?s)person__name">(.*?)</p>', b)
        aff = re.search(r'(?s)person__affil">(.*?)</p>', b)
        talk = re.search(r'(?s)person__talk">(.*?)</p>', b)
        if name:
            rows.append([text_of(name.group(1)),
                         text_of(aff.group(1)) if aff else "",
                         text_of(talk.group(1)) if talk else "—"])
    table(["Name", "Affiliation", "Talk"], rows, widths=[1.7, 2.3, 2.6])

# ==========================================================================
# 9. PROGRAMME
# ==========================================================================
h1("9. Programme")
para("Provisional. The detailed programme, with paper titles, authors and session "
     "chairs, is published on 12 June 2027.", color=GREY, space_after=10)
psrc = read("program.html")
for pm in re.finditer(r'(?s)<div id="day\d" role="tabpanel"[^>]*>(.*?)</div>\s*</div>', psrc):
    panel = pm.group(1)
    hd = re.search(r"(?s)<h3>(.*?)</h3>", panel)
    if hd:
        h2(text_of(hd.group(1)))
    sub = re.search(r'(?s)<p class="muted">(.*?)</p>', panel)
    if sub:
        para(text_of(sub.group(1)), size=9.5, color=GREY, space_after=5)
    rows = []
    for rm in re.finditer(r"(?s)<tr><td class=\"num\">(.*?)</tr>", panel):
        cells = [text_of(c) for c in re.findall(r"(?s)<td[^>]*>(.*?)</td>", "<td>" + rm.group(1))]
        if len(cells) >= 3:
            rows.append(cells[:3])
    table(["Time", "Session", "Venue"], rows, widths=[1.2, 3.6, 1.8])

h2("Presentation guidelines")
gm = re.search(r"(?s)<h2>Presentation guidelines</h2>(.*?)</div>\s*<aside>", psrc)
if gm:
    for tag, txt in blocks(gm.group(1), tags=("h3", "li")):
        if tag == "h3":
            h3(txt)
        else:
            bullet(txt)

# ==========================================================================
# 10. VENUE
# ==========================================================================
h1("10. Venue, Travel and Accommodation")
vsrc = read("venue.html")
h2("Venue address")
for line in ["Sharda School of Computing Science & Engineering",
             "Sharda University",
             "Plot No. 32–34, Knowledge Park III",
             "Greater Noida, Uttar Pradesh — 201310, India",
             "Campus coordinates: 28.4724° N, 77.4834° E"]:
    para(line, space_after=1)

for cap, hd, rows in tables_of(main_of("venue.html")):
    h2("Getting here" if "Arriving by" in (hd[0] if hd else "") else "Accommodation")
    table(hd, rows, widths=[1.7, 2.2, 0.9, 2.0])

h2("Visa and invitation letters")
vm = re.search(r'(?s)<h2>Visa and invitation letters</h2>(.*?)</aside>', vsrc)
if vm:
    for tag, txt in blocks(vm.group(1), tags=("p", "li")):
        if tag == "li":
            numbered(txt) if txt[0].isupper() and len(txt) > 40 else bullet(txt)
        else:
            para(txt)

h2("Places to visit nearby")
am = re.search(r'(?s)id="attractions">(.*?)</section>', vsrc)
if am:
    rows = []
    for cm in re.finditer(r'(?s)<div class="card"><div class="card__body">(.*?)</div></div>', am.group(1)):
        b = cm.group(1)
        d = re.search(r'(?s)trackcard__num[^>]*>(.*?)</p>', b)
        n = re.search(r"(?s)<h3>(.*?)</h3>", b)
        t = re.search(r"(?s)<p>(.*?)</p>", b)
        if n:
            rows.append([text_of(n.group(1)), text_of(d.group(1)) if d else "",
                         text_of(t.group(1)) if t else ""])
    table(["Place", "Distance", "Notes"], rows, widths=[1.3, 1.0, 4.3])

h2("Accessibility")
acc = re.search(r'(?s)id="access">(.*?)</section>', vsrc)
if acc:
    for tag, txt in blocks(acc.group(1), tags=("p", "li")):
        bullet(txt) if tag == "li" else para(txt)

# ==========================================================================
# 11. CONTACT
# ==========================================================================
h1("11. Contact")
csrc = read("contact.html")
h2("Who to write to")
rows = []
for cm in re.finditer(r'(?s)<div class="card"><div class="card__body">\s*<h3>(.*?)</h3>\s*<p>(.*?)</p>\s*<p>(.*?)</p>', csrc):
    rows.append([text_of(cm.group(1)), text_of(cm.group(2)), text_of(cm.group(3))])
table(["Topic", "Use for", "Address"], rows, widths=[1.4, 2.6, 2.6])

h2("Organising secretariat")
rows = []
for pm in re.finditer(r'(?s)<div class="person">(.*?)</div>\s*</div>', csrc):
    b = pm.group(1)
    role = re.search(r'(?s)person__role">(.*?)</p>', b)
    name = re.search(r'(?s)person__name">(.*?)</p>', b)
    aff = re.search(r'(?s)person__affil">(.*?)</p>', b)
    if name:
        rows.append([text_of(name.group(1)),
                     text_of(role.group(1)) if role else "",
                     text_of(aff.group(1)) if aff else ""])
table(["Name", "Role", "Affiliation and contact"], rows, widths=[1.6, 1.5, 3.5])

h2("Postal address")
para("ICNGCI 2027 Secretariat, Sharda School of Computing Science & Engineering, "
     "Sharda University, Plot No. 32–34, Knowledge Park III, Greater Noida, "
     "Uttar Pradesh — 201310, India")

# --------------------------------------------------------------------------
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("wrote", OUT)
print("%.0f KB" % (os.path.getsize(OUT) / 1024))
